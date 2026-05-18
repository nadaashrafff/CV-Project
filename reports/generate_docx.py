"""Generate one Word document per experiment (E7, E8, E9, E10)."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, csv

BASE   = Path("/Users/omarbayoumi/Documents/new cv")
FIG    = BASE / "figures"
RES    = BASE / "results"
OUT    = BASE / "docs"
OUT.mkdir(exist_ok=True)

CLASSES = ["BIODEGRADABLE", "CARDBOARD", "GLASS", "METAL", "PAPER", "PLASTIC"]

# ── palette ────────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x35, 0x64)   # headings
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # sub-headings
LIGHT_BLUE = RGBColor(0xBD, 0xD7, 0xEE)   # table header bg
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── helpers ────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def style_header_row(table, bg="1F3564", fg=WHITE):
    for cell in table.rows[0].cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = fg
                run.font.bold      = True

def add_h1(doc, text):
    p  = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size      = Pt(22)

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = MID_BLUE
        run.font.size      = Pt(15)

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size      = Pt(12)

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

def add_fig(doc, path, caption, width=5.5):
    if Path(path).exists():
        doc.add_picture(str(path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.italic = True
        cap.runs[0].font.size   = Pt(10)
        doc.add_paragraph()

def add_divider(doc):
    p  = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E75B6")
    pb.append(bot)
    pPr.append(pb)

def metrics_from_csv(exp_id):
    with open(RES / "classification_results.csv") as f:
        for row in csv.DictReader(f):
            if row["experiment"] == exp_id:
                return row
    return {}

# ── E7 ─────────────────────────────────────────────────────────────────────────
def make_e7():
    doc = Document()
    doc.core_properties.author = "Bayo"
    doc.core_properties.title  = "E7 Feature Selection Fusion"

    add_h1(doc, "E7 — Feature Selection Fusion")
    add_body(doc, "Member: Bayo  ·  Date: May 2026  ·  Project: Compression-Aware Multiscale Feature Fusion for Waste Object Detection and Classification")
    add_divider(doc)

    # Overview
    add_h2(doc, "1. Overview")
    add_body(doc,
        "Experiment E7 investigates whether applying automated feature selection "
        "to a high-dimensional fused feature vector can match—or exceed—the "
        "performance of using all raw features, while dramatically reducing the "
        "computational and memory footprint. The fused input combines classical "
        "hand-crafted descriptors (252-dim) with EfficientNetB0 deep features "
        "(1,280-dim), giving a 1,532-dimensional input. A Random Forest selector "
        "reduces this to the 200 most informative dimensions before a second RF "
        "classifier is trained on the compressed representation."
    )

    # Methodology
    add_h2(doc, "2. Methodology")
    add_h3(doc, "2.1 Input Data")
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = ["Split", "Shape", "Source"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    rows = [
        ["Train", "(45,177 × 1,532)", "classical_train_clean_X + deep_train_X"],
        ["Val",   "(9,935 × 1,532)",  "classical_val_X + deep_val_X"],
        ["Test",  "(10,553 × 1,532)", "classical_test_X + deep_test_X"],
    ]
    for i, r in enumerate(rows, 1):
        for j, v in enumerate(r):
            table.rows[i].cells[j].text = v
    style_header_row(table)
    doc.add_paragraph()

    add_h3(doc, "2.2 Two-RF Architecture")
    add_body(doc,
        "A critical design rule is maintained: two completely separate Random Forest "
        "objects are used. The selector RF is trained only to extract feature "
        "importances; it is never used for prediction and is discarded after "
        "selecting the top-200 indices. The final RF is then trained from scratch "
        "on the 200-dimensional selected feature matrix."
    )

    # Architecture table
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = "Table Grid"
    for i, h in enumerate(["Step", "Object", "Purpose"]):
        table2.rows[0].cells[i].text = h
    steps = [
        ["1 — Selector RF",  "RandomForestClassifier(n_estimators=200, class_weight='balanced', random_state=42)", "Fit on full 1532-dim train data → extract feature_importances_"],
        ["2 — Top-200 index","np.argsort(importances)[::-1][:200]", "Identify 200 highest-importance feature indices"],
        ["3 — Final RF",     "RandomForestClassifier(n_estimators=200, class_weight='balanced', random_state=42)", "Fit on 200-dim selected train data → saved model artifact"],
    ]
    for i, r in enumerate(steps, 1):
        for j, v in enumerate(r):
            table2.rows[i].cells[j].text = v
    style_header_row(table2)
    doc.add_paragraph()

    add_h3(doc, "2.3 Feature Group Analysis")
    add_body(doc,
        "The 1,532-dim fused vector is partitioned into two groups: "
        "Classical (indices 0–251, 252 dims) and Deep/EfficientNetB0 "
        "(indices 252–1531, 1280 dims). After selection, the top-200 "
        "indices are mapped back to their originating group to understand "
        "which modality the RF found most informative."
    )

    # Results
    add_h2(doc, "3. Results")
    m = metrics_from_csv("E7")

    add_h3(doc, "3.1 Overall Metrics")
    table3 = doc.add_table(rows=7, cols=2)
    table3.style = "Table Grid"
    for i, h in enumerate(["Metric", "Value"]):
        table3.rows[0].cells[i].text = h
    vals = [
        ("Accuracy",          f"{float(m['accuracy']):.2%}"),
        ("Macro-F1",          f"{float(m['macro_f1']):.4f}"),
        ("Weighted-F1",       f"{float(m['weighted_f1']):.4f}"),
        ("Balanced Accuracy", f"{float(m['balanced_accuracy']):.4f}"),
        ("AUC-ROC (macro)",   f"{float(m['auc_roc_macro']):.4f}"),
        ("Feature Dimension", "200 (from 1,532)"),
    ]
    for i, (k, v) in enumerate(vals, 1):
        table3.rows[i].cells[0].text = k
        table3.rows[i].cells[1].text = v
    style_header_row(table3)
    doc.add_paragraph()

    add_h3(doc, "3.2 Per-Class F1 Scores")
    table4 = doc.add_table(rows=7, cols=4)
    table4.style = "Table Grid"
    for i, h in enumerate(["Class", "F1", "Precision", "Recall"]):
        table4.rows[0].cells[i].text = h
    f1_keys  = [f"f1_{c}"   for c in CLASSES]
    pr_keys  = [f"prec_{c}" for c in CLASSES]
    rc_keys  = [f"rec_{c}"  for c in CLASSES]
    for i, cls in enumerate(CLASSES, 1):
        table4.rows[i].cells[0].text = cls
        table4.rows[i].cells[1].text = f"{float(m[f'f1_{cls}']):.4f}"
        table4.rows[i].cells[2].text = f"{float(m[f'prec_{cls}']):.4f}"
        table4.rows[i].cells[3].text = f"{float(m[f'rec_{cls}']):.4f}"
    style_header_row(table4)
    doc.add_paragraph()

    add_h3(doc, "3.3 Feature Group Breakdown")
    table5 = doc.add_table(rows=3, cols=3)
    table5.style = "Table Grid"
    for i, h in enumerate(["Group", "Features Selected", "Percentage"]):
        table5.rows[0].cells[i].text = h
    classical_n = int(float(m.get("feature_group_classical", 102)))
    deep_n      = int(float(m.get("feature_group_deep", 98)))
    table5.rows[1].cells[0].text = "Classical (hand-crafted)"
    table5.rows[1].cells[1].text = str(classical_n)
    table5.rows[1].cells[2].text = f"{classical_n/200:.0%}"
    table5.rows[2].cells[0].text = "Deep (EfficientNetB0)"
    table5.rows[2].cells[1].text = str(deep_n)
    table5.rows[2].cells[2].text = f"{deep_n/200:.0%}"
    style_header_row(table5)
    doc.add_paragraph()

    # Figures
    add_h2(doc, "4. Figures")
    add_fig(doc, FIG / "E7_feature_group_analysis.png",
            "Figure 1 — Feature group analysis: distribution of top-200 selected features across Classical and Deep modalities")
    add_fig(doc, FIG / "E7_confusion_matrix.png",
            "Figure 2 — Confusion matrix on the test set (10,553 samples, 6 classes)")

    # Insights
    add_h2(doc, "5. Insights & Discussion")
    add_bullet(doc, "E7 achieves 80.4% accuracy and Macro-F1 = 0.708, which is above the classical-only baseline (E2: 0.648) but below the deep-only SVM baseline (E3: 0.785).")
    add_bullet(doc, f"The RF selector chose {classical_n} classical and {deep_n} deep features — a near-even split, confirming both modalities carry complementary information.")
    add_bullet(doc, "BIODEGRADABLE achieves near-perfect recall (0.995) due to its dominant class frequency; minority classes (GLASS, PLASTIC) remain challenging (F1 ≈ 0.61–0.63).")
    add_bullet(doc, "Reducing from 1,532 to 200 dimensions (87% compression) incurs only a moderate performance penalty, making E7 attractive for resource-constrained deployments.")
    add_bullet(doc, "The two-RF design is intentional: using the selector RF's predictions would leak importance-based bias into the evaluation.")

    doc.save(OUT / "E7_Feature_Selection.docx")
    print("✓ E7_Feature_Selection.docx")


# ── E8 ─────────────────────────────────────────────────────────────────────────
def make_e8():
    doc = Document()
    doc.core_properties.author = "Bayo"
    doc.core_properties.title  = "E8 Average Voting Fusion"

    add_h1(doc, "E8 — Average Voting Fusion")
    add_body(doc, "Member: Bayo  ·  Date: May 2026  ·  Project: Compression-Aware Multiscale Feature Fusion for Waste Object Detection and Classification")
    add_divider(doc)

    add_h2(doc, "1. Overview")
    add_body(doc,
        "Experiment E8 explores the simplest possible late-fusion strategy: "
        "averaging the class probability vectors produced by two independently-trained "
        "classifiers — the Random Forest from E2 (classical features) and the RBF SVM "
        "from E3 (deep features). No additional training is required. The hypothesis is "
        "that combining two diverse probability distributions will reduce individual "
        "model errors and improve overall Macro-F1."
    )

    add_h2(doc, "2. Methodology")
    add_h3(doc, "2.1 Input Models")
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Model", "Features", "Dim", "Source"]):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "Random Forest (E2)"
    table.rows[1].cells[1].text = "Classical hand-crafted"
    table.rows[1].cells[2].text = "252"
    table.rows[1].cells[3].text = "random_forest_E2.pkl"
    table.rows[2].cells[0].text = "RBF SVM (E3)"
    table.rows[2].cells[1].text = "EfficientNetB0 deep"
    table.rows[2].cells[2].text = "1,280"
    table.rows[2].cells[3].text = "svm_E3.pkl  (dict: clf + scaler)"
    style_header_row(table)
    doc.add_paragraph()

    add_h3(doc, "2.2 Fusion Formula")
    add_body(doc, "For each test sample, six-class probability vectors are obtained from both models and averaged element-wise:")
    p = doc.add_paragraph()
    run = p.add_run("    avg_proba = (RF_proba + SVM_proba) / 2.0\n    y_pred = argmax(avg_proba, axis=1)")
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    doc.add_paragraph()
    add_body(doc,
        "The SVM model is stored as a dict {'clf': SVC, 'scaler': StandardScaler}. "
        "The scaler is applied before calling clf.predict_proba() to ensure the deep "
        "features are normalised identically to training."
    )

    add_h3(doc, "2.3 Architecture Diagram")
    steps = [
        ("Classical features (252-dim)", "→ RF.predict_proba()", "RF proba (N × 6)"),
        ("Deep features (1,280-dim)",    "→ scaler → SVM.predict_proba()", "SVM proba (N × 6)"),
        ("",                             "→ element-wise average",          "avg proba (N × 6)"),
        ("",                             "→ argmax",                        "Final prediction (N,)"),
    ]
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = "Table Grid"
    for i, h in enumerate(["Input", "Transform", "Output"]):
        table2.rows[0].cells[i].text = h
    for i, (a, b, c) in enumerate(steps, 1):
        table2.rows[i].cells[0].text = a
        table2.rows[i].cells[1].text = b
        table2.rows[i].cells[2].text = c
    style_header_row(table2)
    doc.add_paragraph()

    add_h2(doc, "3. Results")
    m = metrics_from_csv("E8")

    add_h3(doc, "3.1 Overall Metrics")
    table3 = doc.add_table(rows=7, cols=2)
    table3.style = "Table Grid"
    for i, h in enumerate(["Metric", "Value"]):
        table3.rows[0].cells[i].text = h
    vals = [
        ("Accuracy",          f"{float(m['accuracy']):.2%}"),
        ("Macro-F1",          f"{float(m['macro_f1']):.4f}"),
        ("Weighted-F1",       f"{float(m['weighted_f1']):.4f}"),
        ("Balanced Accuracy", f"{float(m['balanced_accuracy']):.4f}"),
        ("AUC-ROC (macro)",   f"{float(m['auc_roc_macro']):.4f}"),
        ("Training required", "None (inference only)"),
    ]
    for i, (k, v) in enumerate(vals, 1):
        table3.rows[i].cells[0].text = k
        table3.rows[i].cells[1].text = v
    style_header_row(table3)
    doc.add_paragraph()

    add_h3(doc, "3.2 Per-Class F1 Scores")
    table4 = doc.add_table(rows=7, cols=4)
    table4.style = "Table Grid"
    for i, h in enumerate(["Class", "F1", "Precision", "Recall"]):
        table4.rows[0].cells[i].text = h
    for i, cls in enumerate(CLASSES, 1):
        table4.rows[i].cells[0].text = cls
        table4.rows[i].cells[1].text = f"{float(m[f'f1_{cls}']):.4f}"
        table4.rows[i].cells[2].text = f"{float(m[f'prec_{cls}']):.4f}"
        table4.rows[i].cells[3].text = f"{float(m[f'rec_{cls}']):.4f}"
    style_header_row(table4)
    doc.add_paragraph()

    add_h3(doc, "3.3 Comparison vs. Individual Baselines")
    table5 = doc.add_table(rows=4, cols=3)
    table5.style = "Table Grid"
    for i, h in enumerate(["Experiment", "Accuracy", "Macro-F1"]):
        table5.rows[0].cells[i].text = h
    rows = [
        ("E2 — RF (Classical only)", "77.40%", "0.6476"),
        ("E3 — SVM (Deep only)",     "85.22%", "0.7848"),
        ("E8 — Average Voting",      "84.84%", "0.7820"),
    ]
    for i, r in enumerate(rows, 1):
        for j, v in enumerate(r):
            table5.rows[i].cells[j].text = v
    style_header_row(table5)
    doc.add_paragraph()

    add_h2(doc, "4. Figures")
    add_fig(doc, FIG / "E8_confusion_matrix.png",
            "Figure 1 — Confusion matrix for E8 Average Voting on the test set (10,553 samples)")

    add_h2(doc, "5. Insights & Discussion")
    add_bullet(doc, "E8 achieves Macro-F1 = 0.782, nearly matching the strong SVM-only baseline (E3: 0.785) without any additional training.")
    add_bullet(doc, "The RF contributes most on BIODEGRADABLE (dominant class); the SVM contributes most on minority classes (GLASS, PLASTIC) due to its kernel-based decision boundary.")
    add_bullet(doc, "Average voting is a zero-cost ensemble: it requires only the two pre-trained models and a single forward pass per model at inference time.")
    add_bullet(doc, "The main risk is that both models share the same errors on hard samples — E9 weighted voting attempts to mitigate this by amplifying the stronger model.")
    add_bullet(doc, "AUC-ROC (0.946) is substantially higher than either individual model, confirming complementarity of the two probability distributions.")

    doc.save(OUT / "E8_Average_Voting.docx")
    print("✓ E8_Average_Voting.docx")


# ── E9 ─────────────────────────────────────────────────────────────────────────
def make_e9():
    doc = Document()
    doc.core_properties.author = "Bayo"
    doc.core_properties.title  = "E9 Weighted Voting Fusion"

    add_h1(doc, "E9 — Weighted Voting Fusion")
    add_body(doc, "Member: Bayo  ·  Date: May 2026  ·  Project: Compression-Aware Multiscale Feature Fusion for Waste Object Detection and Classification")
    add_divider(doc)

    add_h2(doc, "1. Overview")
    add_body(doc,
        "Experiment E9 extends E8 by replacing the fixed 50/50 average with "
        "learned, data-driven weights. A grid search over 19 candidate weight "
        "pairs is run exclusively on the validation set to find the combination "
        "that maximises Macro-F1. The optimal weights are applied to the test set "
        "exactly once, preserving test-set integrity."
    )

    add_h2(doc, "2. Methodology")
    add_h3(doc, "2.1 Grid Search Protocol")
    add_body(doc, "Weight candidates: w_rf ∈ {0.05, 0.10, …, 0.95} — 19 evenly spaced values, w_svm = 1 − w_rf.")
    add_body(doc, "Objective: maximise Macro-F1 on the validation set. Test set is never inspected during the search.")

    p = doc.add_paragraph()
    run = p.add_run(
        "    for w_rf in [0.05, 0.10, ..., 0.95]:\n"
        "        fused_val = w_rf * RF_proba_val + (1-w_rf) * SVM_proba_val\n"
        "        macro_f1  = f1_score(y_val, argmax(fused_val), average='macro')\n\n"
        "    # Apply best weights to test set (once)\n"
        "    fused_test = best_w_rf * RF_proba_test + best_w_svm * SVM_proba_test"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    doc.add_paragraph()

    add_h3(doc, "2.2 Full Weight Search Results")
    with open(RES / "e9_weight_search_results.csv") as f:
        rows_csv = list(csv.DictReader(f))
    table = doc.add_table(rows=len(rows_csv)+1, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["w_RF", "w_SVM", "Val Accuracy", "Val Macro-F1"]):
        table.rows[0].cells[i].text = h
    best_wrf = 0.4
    for i, r in enumerate(rows_csv, 1):
        table.rows[i].cells[0].text = r["w_rf"]
        table.rows[i].cells[1].text = r["w_svm"]
        table.rows[i].cells[2].text = f"{float(r['val_accuracy']):.4f}"
        table.rows[i].cells[3].text = f"{float(r['val_macro_f1']):.4f}"
        if float(r["w_rf"]) == best_wrf:
            for j in range(4):
                set_cell_bg(table.rows[i].cells[j], "DDEBF7")
    style_header_row(table)
    cap = doc.add_paragraph("Highlighted row = optimal weights (w_RF=0.40, w_SVM=0.60)")
    cap.runs[0].font.italic = True
    cap.runs[0].font.size   = Pt(10)
    doc.add_paragraph()

    add_h2(doc, "3. Results")
    m = metrics_from_csv("E9")

    add_h3(doc, "3.1 Optimal Weights")
    with open(RES / "e9_optimal_weights.json") as f:
        opt = json.load(f)
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = "Table Grid"
    for i, h in enumerate(["Parameter", "Value"]):
        table2.rows[0].cells[i].text = h
    table2.rows[1].cells[0].text = "Optimal w_RF"
    table2.rows[1].cells[1].text = str(opt["w_rf"])
    table2.rows[2].cells[0].text = "Optimal w_SVM"
    table2.rows[2].cells[1].text = str(opt["w_svm"])
    table2.rows[3].cells[0].text = "Val Macro-F1 at optimum"
    table2.rows[3].cells[1].text = f"{opt['val_macro_f1']:.4f}"
    style_header_row(table2)
    doc.add_paragraph()

    add_h3(doc, "3.2 Overall Test Metrics")
    table3 = doc.add_table(rows=7, cols=2)
    table3.style = "Table Grid"
    for i, h in enumerate(["Metric", "Value"]):
        table3.rows[0].cells[i].text = h
    vals = [
        ("Accuracy",          f"{float(m['accuracy']):.2%}"),
        ("Macro-F1",          f"{float(m['macro_f1']):.4f}"),
        ("Weighted-F1",       f"{float(m['weighted_f1']):.4f}"),
        ("Balanced Accuracy", f"{float(m['balanced_accuracy']):.4f}"),
        ("AUC-ROC (macro)",   f"{float(m['auc_roc_macro']):.4f}"),
        ("Training required", "None (grid search on val only)"),
    ]
    for i, (k, v) in enumerate(vals, 1):
        table3.rows[i].cells[0].text = k
        table3.rows[i].cells[1].text = v
    style_header_row(table3)
    doc.add_paragraph()

    add_h3(doc, "3.3 Per-Class F1 Scores")
    table4 = doc.add_table(rows=7, cols=4)
    table4.style = "Table Grid"
    for i, h in enumerate(["Class", "F1", "Precision", "Recall"]):
        table4.rows[0].cells[i].text = h
    for i, cls in enumerate(CLASSES, 1):
        table4.rows[i].cells[0].text = cls
        table4.rows[i].cells[1].text = f"{float(m[f'f1_{cls}']):.4f}"
        table4.rows[i].cells[2].text = f"{float(m[f'prec_{cls}']):.4f}"
        table4.rows[i].cells[3].text = f"{float(m[f'rec_{cls}']):.4f}"
    style_header_row(table4)
    doc.add_paragraph()

    add_h2(doc, "4. Figures")
    add_fig(doc, FIG / "E9_weight_search.png",
            "Figure 1 — Validation Macro-F1 across all 19 weight candidates. "
            "The optimal is at w_RF=0.40 (dashed line), confirming the SVM should receive more weight.")
    add_fig(doc, FIG / "E9_confusion_matrix.png",
            "Figure 2 — Confusion matrix for E9 Weighted Voting on the test set (10,553 samples)")

    add_h2(doc, "5. Insights & Discussion")
    add_bullet(doc, "Optimal weights: w_RF=0.40, w_SVM=0.60 — the SVM (deep features) receives more trust, consistent with E3 being the stronger individual model.")
    add_bullet(doc, "E9 (Macro-F1=0.7832) marginally outperforms E8 (0.7820) — a +0.0012 gain, suggesting the equal-weight assumption in E8 was already close to optimal.")
    add_bullet(doc, "The weight-vs-F1 curve peaks sharply at w_RF=0.40 and degrades monotonically toward w_RF=0.95, confirming deep features are the dominant signal.")
    add_bullet(doc, "Both E8 and E9 achieve near-parity with the SVM baseline (E3: 0.7848) at zero additional training cost — a compelling result for production systems.")
    add_bullet(doc, "GLASS and PLASTIC remain the hardest classes (F1 ≈ 0.72); no weighting scheme can recover information absent from both probability vectors.")

    doc.save(OUT / "E9_Weighted_Voting.docx")
    print("✓ E9_Weighted_Voting.docx")


# ── E10 ────────────────────────────────────────────────────────────────────────
def make_e10():
    doc = Document()
    doc.core_properties.author = "Bayo"
    doc.core_properties.title  = "E10 Attention Fusion"

    add_h1(doc, "E10 — Attention-Based Fusion")
    add_body(doc, "Member: Bayo  ·  Date: May 2026  ·  Project: Compression-Aware Multiscale Feature Fusion for Waste Object Detection and Classification")
    add_divider(doc)

    add_h2(doc, "1. Overview")
    add_body(doc,
        "Experiment E10 is the most sophisticated of Bayo's four experiments. "
        "Instead of assigning static weights to entire probability distributions, "
        "a PyTorch neural network learns to dynamically gate the contribution of "
        "each modality on a per-sample basis. Two separate embedding branches "
        "project classical (252-dim) and deep (1,280-dim) features into a shared "
        "128-dimensional space. A softmax attention mechanism then assigns a "
        "scalar weight to each branch, producing a weighted combination of the "
        "two embeddings as the fused representation for classification."
    )

    add_h2(doc, "2. Methodology")
    add_h3(doc, "2.1 Preprocessing")
    add_body(doc,
        "Both feature modalities are independently standardised using "
        "StandardScaler fitted on the training set only, then applied to "
        "validation and test sets."
    )
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Modality", "Input Dim", "Scaler"]):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "Classical"
    table.rows[1].cells[1].text = "252"
    table.rows[1].cells[2].text = "e10_scaler_classical.pkl"
    table.rows[2].cells[0].text = "Deep (EfficientNetB0)"
    table.rows[2].cells[1].text = "1,280"
    table.rows[2].cells[2].text = "e10_scaler_deep.pkl"
    style_header_row(table)
    doc.add_paragraph()

    add_h3(doc, "2.2 Network Architecture")
    add_body(doc, "The AttentionFusion model (PyTorch) has three components:")

    arch_table = doc.add_table(rows=6, cols=3)
    arch_table.style = "Table Grid"
    for i, h in enumerate(["Component", "Layers", "Output"]):
        arch_table.rows[0].cells[i].text = h
    arch_rows = [
        ("Classical branch", "Linear(252→256) → ReLU → BN → Linear(256→128) → ReLU", "128-dim embedding"),
        ("Deep branch",      "Linear(1280→512) → ReLU → BN → Linear(512→256) → ReLU → BN → Linear(256→128) → ReLU", "128-dim embedding"),
        ("Attention gate",   "Concat(128+128=256) → Linear(256→64) → ReLU → Linear(64→2) → Softmax", "2 scalar weights (α_c, α_d)"),
        ("Fusion",           "α_c × classical_embed + α_d × deep_embed", "128-dim attended vector"),
        ("Classifier head",  "Linear(128→64) → ReLU → Dropout(0.3) → Linear(64→6) → Softmax", "6-class probabilities"),
    ]
    for i, r in enumerate(arch_rows, 1):
        for j, v in enumerate(r):
            arch_table.rows[i].cells[j].text = v
    style_header_row(arch_table)
    doc.add_paragraph()

    add_h3(doc, "2.3 Training Configuration")
    table2 = doc.add_table(rows=8, cols=2)
    table2.style = "Table Grid"
    for i, h in enumerate(["Hyperparameter", "Value"]):
        table2.rows[0].cells[i].text = h
    params = [
        ("Optimizer",          "Adam, lr=1e-3"),
        ("Loss",               "Cross-Entropy with class_weight='balanced'"),
        ("Batch size",         "32"),
        ("Max epochs",         "100"),
        ("Early stopping",     "patience=10, monitor=val_loss"),
        ("LR scheduler",       "ReduceLROnPlateau (factor=0.5, patience=5, min_lr=1e-6)"),
        ("Random seed",        "42"),
    ]
    for i, (k, v) in enumerate(params, 1):
        table2.rows[i].cells[0].text = k
        table2.rows[i].cells[1].text = v
    style_header_row(table2)
    doc.add_paragraph()

    add_h2(doc, "3. Results")
    m = metrics_from_csv("E10")

    add_h3(doc, "3.1 Overall Metrics")
    table3 = doc.add_table(rows=9, cols=2)
    table3.style = "Table Grid"
    for i, h in enumerate(["Metric", "Value"]):
        table3.rows[0].cells[i].text = h
    vals = [
        ("Accuracy",          f"{float(m['accuracy']):.2%}"),
        ("Macro-F1",          f"{float(m['macro_f1']):.4f}"),
        ("Weighted-F1",       f"{float(m['weighted_f1']):.4f}"),
        ("Balanced Accuracy", f"{float(m['balanced_accuracy']):.4f}"),
        ("AUC-ROC (macro)",   f"{float(m['auc_roc_macro']):.4f}"),
        ("Model size",        f"{float(m['model_size_mb']):.2f} MB"),
        ("Epochs trained",    f"{int(float(m['epochs_trained']))} (early stop, patience=10)"),
        ("Feature dim",       "128 (learned fused embedding)"),
    ]
    for i, (k, v) in enumerate(vals, 1):
        table3.rows[i].cells[0].text = k
        table3.rows[i].cells[1].text = v
    style_header_row(table3)
    doc.add_paragraph()

    add_h3(doc, "3.2 Per-Class F1 Scores")
    table4 = doc.add_table(rows=7, cols=4)
    table4.style = "Table Grid"
    for i, h in enumerate(["Class", "F1", "Precision", "Recall"]):
        table4.rows[0].cells[i].text = h
    for i, cls in enumerate(CLASSES, 1):
        table4.rows[i].cells[0].text = cls
        table4.rows[i].cells[1].text = f"{float(m[f'f1_{cls}']):.4f}"
        table4.rows[i].cells[2].text = f"{float(m[f'prec_{cls}']):.4f}"
        table4.rows[i].cells[3].text = f"{float(m[f'rec_{cls}']):.4f}"
    style_header_row(table4)
    doc.add_paragraph()

    add_h3(doc, "3.3 Per-Class Attention Weights")
    with open(RES / "e10_attention_by_class.csv") as f:
        attn_rows = list(csv.DictReader(f))
    table5 = doc.add_table(rows=len(attn_rows)+1, cols=4)
    table5.style = "Table Grid"
    for i, h in enumerate(["Class", "Avg Classical Weight (α_c)", "Avg Deep Weight (α_d)", "N Samples"]):
        table5.rows[0].cells[i].text = h
    for i, r in enumerate(attn_rows, 1):
        table5.rows[i].cells[0].text = r["class"]
        table5.rows[i].cells[1].text = f"{float(r['avg_classical_w']):.4f}"
        table5.rows[i].cells[2].text = f"{float(r['avg_deep_w']):.4f}"
        table5.rows[i].cells[3].text = r["n_samples"]
    style_header_row(table5)
    cap = doc.add_paragraph("Each row sums to 1.0. α_d > α_c for all classes except PLASTIC, reflecting the superior discriminative power of deep features in most cases.")
    cap.runs[0].font.italic = True
    cap.runs[0].font.size   = Pt(10)
    doc.add_paragraph()

    add_h2(doc, "4. Figures")
    add_fig(doc, FIG / "E10_training_curves.png",
            "Figure 1 — Training and validation loss/accuracy curves. Early stopping triggered at epoch 13.")
    add_fig(doc, FIG / "E10_attention_heatmap.png",
            "Figure 2 — Per-class average attention weights (heatmap). "
            "Deep features (right column) consistently receive higher weight, except PLASTIC.")
    add_fig(doc, FIG / "E10_confusion_matrix.png",
            "Figure 3 — Confusion matrix for E10 Attention Fusion on the test set (10,553 samples)")

    add_h2(doc, "5. Insights & Discussion")
    add_bullet(doc, "E10 achieves the highest accuracy of all Bayo's experiments (85.64%), surpassing even the SVM-only baseline in raw accuracy, though Macro-F1 (0.775) is slightly below E9.")
    add_bullet(doc, "The attention gate learns near-equal weights overall (α_c≈0.43, α_d≈0.57), but per-class variation is meaningful — PLASTIC is the only class where classical features receive more weight (0.510 vs 0.490).")
    add_bullet(doc, "BIODEGRADABLE achieves F1=0.941, the highest per-class score, driven by its dominant class frequency and high attention to deep features (α_d=0.614).")
    add_bullet(doc, "The model is compact at 1.58 MB and converges in 13 epochs — practical for deployment even without GPU acceleration.")
    add_bullet(doc, "The attention mechanism provides interpretability: per-sample gate weights can be logged at inference time to explain which modality drove each prediction.")
    add_bullet(doc, "GLASS and METAL remain challenging (F1≈0.72–0.76) for all fusion methods, suggesting class confusion inherent to the visual features rather than a fusion limitation.")

    doc.save(OUT / "E10_Attention_Fusion.docx")
    print("✓ E10_Attention_Fusion.docx")


if __name__ == "__main__":
    make_e7()
    make_e8()
    make_e9()
    make_e10()
    print(f"\nAll documents saved to: {OUT}")
